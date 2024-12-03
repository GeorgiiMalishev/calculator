import io
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches
from fpdf import FPDF



def convert_to_html(workbook):
    """
    Преобразует данные из Excel в HTML-таблицу с поддержкой объединенных ячеек.

    Параметры:
        workbook: Объект Excel-книги.

    Возвращает:
        HTML-строка с таблицей.
    """
    sheet = workbook.active

    # Создаем матрицу для отслеживания обработанных ячеек
    max_row = sheet.max_row
    max_col = sheet.max_column
    processed_cells = [[False] * max_col for _ in range(max_row)]

    html = "<table style='width: 100%; border-collapse: collapse;'>\n"

    for row_idx, row in enumerate(sheet.iter_rows(), 1):
        html += "  <tr>\n"
        for col_idx, cell in enumerate(row, 1):
            # Пропускаем уже обработанные ячейки
            if processed_cells[row_idx - 1][col_idx - 1]:
                continue

            # Проверяем, является ли ячейка частью объединенного диапазона
            colspan = 1
            rowspan = 1
            for merged_range in sheet.merged_cells.ranges:
                if cell.coordinate in merged_range:
                    min_col, min_row, max_col, max_row = merged_range.bounds
                    colspan = max_col - min_col + 1
                    rowspan = max_row - min_row + 1

                    # Помечаем все ячейки в объединенном диапазоне как обработанные
                    for r in range(min_row, max_row + 1):
                        for c in range(min_col, max_col + 1):
                            processed_cells[r - 1][c - 1] = True

                    break

            value = cell.value if cell.value is not None else ""

            # Добавляем стили для центрирования и объединения
            style = 'border: 1px solid #ddd; padding: 8px; text-align: center; vertical-align: middle;'
            span_attrs = ''
            if colspan > 1:
                span_attrs += f' colspan="{colspan}"'
            if rowspan > 1:
                span_attrs += f' rowspan="{rowspan}"'

            html += f"    <td style='{style}'{span_attrs}>{value}</td>\n"

        html += "  </tr>\n"
    html += "</table>"
    return html


def convert_excel_to_word(excel):
    """
    Преобразует данные из Excel в Word-документ с поддержкой объединенных ячеек
    в альбомной ориентации.

    Параметры:
        excel: Объект Excel-книги.

    Возвращает:
        Word-документ в байтовом формате.
    """
    doc = Document()

    # Устанавливаем альбомную ориентацию
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    # Устанавливаем размеры страницы для альбомной ориентации
    section.page_width = Inches(11.69)  # A4 высота становится шириной
    section.page_height = Inches(8.27)  # A4 ширина становится высотой
    # Устанавливаем поля
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    doc.add_heading('Экспортированная таблица', level=1)

    sheet = excel.active
    table = doc.add_table(rows=sheet.max_row, cols=sheet.max_column)
    table.style = 'Table Grid'

    # Создаем матрицу для отслеживания объединенных ячеек
    merged_cells = {}
    for merged_range in sheet.merged_cells.ranges:
        min_col, min_row, max_col, max_row = merged_range.bounds
        for row in range(min_row, max_row + 1):
            for col in range(min_col, max_col + 1):
                merged_cells[(row, col)] = (min_row, min_col, max_row - min_row + 1, max_col - min_col + 1)

    # Заполняем таблицу с учетом объединенных ячеек
    for row_idx, row in enumerate(sheet.rows, 1):
        for col_idx, cell in enumerate(row, 1):
            word_cell = table.cell(row_idx - 1, col_idx - 1)

            if (row_idx, col_idx) in merged_cells:
                start_row, start_col, height, width = merged_cells[(row_idx, col_idx)]
                if row_idx == start_row and col_idx == start_col:
                    if height > 1:
                        word_cell.merge(table.cell(row_idx + height - 2, col_idx - 1))
                    if width > 1:
                        word_cell.merge(table.cell(row_idx - 1, col_idx + width - 2))

                    value = sheet.cell(start_row, start_col).value
                    word_cell.text = str(value) if value is not None else ""
            else:
                value = cell.value
                word_cell.text = str(value) if value is not None else ""

            # Применяем форматирование
            for paragraph in word_cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                paragraph.style.font.size = Pt(10)

    doc_output = io.BytesIO()
    doc.save(doc_output)
    doc_output.seek(0)
    return doc_output



def convert_excel_to_pdf(excel):
    """
    Конвертирует Excel в PDF с корректной обработкой объединенных ячеек
    """
    pdf = FPDF(orientation='L')
    pdf.add_page()
    pdf.add_font('FreeSans', '', 'static/font/FreeSans.ttf', uni=True)
    pdf.set_font('FreeSans', '', 8)

    sheet = excel.active
    margin = 10
    page_width = pdf.w - 2 * margin
    max_col = sheet.max_column
    col_width = page_width / max_col
    base_row_height = 8

    # Вычисляем X-координату для каждого столбца
    col_x_positions = {col: margin + (col - 1) * col_width for col in range(1, max_col + 1)}

    # Создаем структуру данных для хранения информации о merged cells
    merged_cells_map = {}
    for merged_range in sheet.merged_cells.ranges:
        min_col, min_row, max_col, max_row = merged_range.bounds
        width = (max_col - min_col + 1) * col_width
        for row in range(min_row, max_row + 1):
            for col in range(min_col, max_col + 1):
                merged_cells_map[(row, col)] = {
                    'start': (min_row, min_col),
                    'end': (max_row, max_col),
                    'width': width,
                    'x_position': col_x_positions[min_col],
                    'min_col': min_col,
                    'max_col': max_col,
                    'is_main': (row, col) == (min_row, min_col)
                }

    def get_text_dimensions(text, width):
        """Вычисляет размеры текста с учетом переносов"""
        if text is None:
            return 0
        lines = pdf.multi_cell(width - 2, base_row_height, str(text), split_only=True)
        return len(lines) * base_row_height + 4

    # Вычисляем высоты строк
    row_heights = {}
    max_row = sheet.max_row

    for row_idx in range(1, max_row + 1):
        row = sheet.iter_rows(min_row=row_idx, max_row=row_idx, min_col=1, max_col=max_col).__next__()
        
        # Пропускаем пустые строки
        if not any(cell.value is not None for cell in row):
            row_heights[row_idx] = base_row_height
            continue

        max_height = base_row_height
        col_idx = 1
        
        while col_idx <= max_col:
            cell = sheet.cell(row_idx, col_idx)
            
            if (row_idx, col_idx) in merged_cells_map:
                merge_info = merged_cells_map[(row_idx, col_idx)]
                
                if merge_info['is_main']:
                    # Вычисляем высоту текста для объединенной ячейки
                    text_height = get_text_dimensions(cell.value, merge_info['width'])
                    merge_height = max(text_height, base_row_height)
                    
                    # Распределяем высоту между строками объединенной ячейки
                    rows_count = merge_info['end'][0] - merge_info['start'][0] + 1
                    distributed_height = max(merge_height / rows_count, base_row_height)
                    
                    for r in range(merge_info['start'][0], merge_info['end'][0] + 1):
                        row_heights[r] = max(row_heights.get(r, base_row_height), distributed_height)
                    
                    col_idx = merge_info['max_col'] + 1
                    continue
            
            # Обычная ячейка
            text_height = get_text_dimensions(cell.value, col_width)
            max_height = max(max_height, text_height)
            col_idx += 1
        
        # Устанавливаем высоту для обычных ячеек
        if row_idx not in row_heights:
            row_heights[row_idx] = max(max_height, base_row_height)

    # Отрисовка таблицы
    current_y = margin
    processed_cells = set()
    max_col = sheet.max_column
    for row_idx in range(1, max_row + 1):
        row = sheet.iter_rows(min_row=row_idx, max_row=row_idx, min_col=1, max_col=max_col).__next__()
        
        if not any(cell.value is not None for cell in row):
            continue

        row_height = row_heights[row_idx]

        # Проверяем переход на новую страницу
        if current_y + row_height > pdf.h - margin:
            pdf.add_page()
            current_y = margin

        col_idx = 1
        while col_idx <= max_col:
            if (row_idx, col_idx) in processed_cells:
                col_idx += 1
                continue

            if (row_idx, col_idx) in merged_cells_map:
                merge_info = merged_cells_map[(row_idx, col_idx)]
                
                if merge_info['is_main']:
                    # Вычисляем общую высоту объединенной ячейки
                    total_height = sum(row_heights[r] for r in range(
                        merge_info['start'][0],
                        merge_info['end'][0] + 1
                    ))

                    cell = sheet.cell(row_idx, col_idx)
                    x_position = merge_info['x_position']
                    
                    # Отрисовка объединенной ячейки
                    pdf.set_xy(x_position, current_y)
                    pdf.cell(merge_info['width'], total_height, '', 1)

                    # Отрисовка текста в объединенной ячейке
                    if cell.value is not None:
                        text = str(cell.value)
                        effective_width = merge_info['width'] - 2
                        lines = pdf.multi_cell(effective_width, base_row_height, text, split_only=True)
                        text_height = len(lines) * base_row_height
                        text_y = current_y + (total_height - text_height) / 2
                        pdf.set_xy(x_position + 1, text_y)
                        pdf.multi_cell(effective_width, base_row_height, text, align='C')

                    # Отмечаем все ячейки в диапазоне как обработанные
                    for r in range(merge_info['start'][0], merge_info['end'][0] + 1):
                        for c in range(merge_info['start'][1], merge_info['end'][1] + 1):
                            processed_cells.add((r, c))

                    col_idx = merge_info['max_col'] + 1
                    continue
                else:
                    col_idx += 1
                    continue

            # Отрисовка обычной ячейки
            cell = sheet.cell(row_idx, col_idx)
            x_position = col_x_positions[col_idx]
            
            pdf.set_xy(x_position, current_y)
            pdf.cell(col_width, row_height, '', 1)

            if cell.value is not None:
                text = str(cell.value)
                effective_width = col_width - 2
                lines = pdf.multi_cell(effective_width, base_row_height, text, split_only=True)
                text_height = len(lines) * base_row_height
                text_y = current_y + (row_height - text_height) / 2
                pdf.set_xy(x_position + 1, text_y)
                pdf.multi_cell(effective_width, base_row_height, text, align='C')

            processed_cells.add((row_idx, col_idx))
            col_idx += 1

        current_y += row_height

    return pdf.output(dest='S').encode('latin1')