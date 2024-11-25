import io
import json
from django.shortcuts import render
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from fpdf import FPDF
from calculatorApp import calculator1, resultToExcel1
from calculatorApp.calculator1 import Calculator1


@csrf_exempt
def calculate_fact(request):
    """
    Обрабатывает POST-запросы для расчёта фактических данных.
    Принимает JSON-данные, обрабатывает их с помощью Calculator1 и возвращает результат в виде HTML.

    Параметры:
        request: Объект запроса Django, содержащий JSON-данные.

    Возвращает:
        JsonResponse с HTML-результатом или сообщением об ошибке.
    """
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            data = {key: int(value) for key, value in data.items()}
            calculator = Calculator1(data)
            results = calculator.calculate()
            excel = resultToExcel1.fact(results)
            html_content = convert_to_html(excel)
            context = {'html_content': html_content}
            return JsonResponse(context)
        except json.JSONDecodeError:
            return JsonResponse({'status': 'error', 'message': 'Некорректный формат данных'}, status=400)
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)
    return JsonResponse({'status': 'error', 'message': 'Метод не поддерживается'}, status=405)


@csrf_exempt
def calculate_plan(request):
    """
    Обрабатывает POST-запросы для расчёта плановых данных.
    Принимает JSON-данные, обрабатывает их с помощью Calculator1 и возвращает результат в виде HTML.

    Параметры:
        request: Объект запроса Django, содержащий JSON-данные.

    Возвращает:
        JsonResponse с HTML-результатом или сообщением об ошибке.
    """
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            data = {key: int(value) for key, value in data.items()}
            calculator = Calculator1(data)
            results = calculator.calculate()
            excel = resultToExcel1.plan(results)
            html_content = convert_to_html(excel)
            context = {'html_content': html_content}
            return JsonResponse(context)
        except json.JSONDecodeError:
            return JsonResponse({'status': 'error', 'message': 'Некорректный формат данных'}, status=400)
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)
    return JsonResponse({'status': 'error', 'message': 'Метод не поддерживается'}, status=405)


def calculator_view(request, calculator_name):
    """
    Рендерит HTML-страницу для заданного калькулятора.

    Параметры:
        request: Объект запроса Django.
        calculator_name: Имя калькулятора для отображения соответствующей страницы.

    Возвращает:
        HttpResponse с HTML-страницей.
    """
    return render(request, f'calculator{calculator_name}.html')


def convert_to_html(workbook):
    """
    Преобразует данные из Excel в HTML-таблицу.

    Параметры:
        workbook: Объект Excel-книги.

    Возвращает:
        HTML-строка с таблицей.
    """
    sheet = workbook.active
    html = "<table style='width: 100%; border-collapse: collapse;'>\n"
    for row in sheet.iter_rows():
        html += "  <tr>\n"
        for cell in row:
            value = cell.value if cell.value is not None else ""
            for merged_range in sheet.merged_cells.ranges:
                if cell.coordinate in merged_range:
                    top_left_cell = sheet[merged_range.coord.split(":")[0]]
                    value = top_left_cell.value
                    break
            html += f"    <td style='border: 1px solid #ddd; padding: 8px;'>{value}</td>\n"
        html += "  </tr>\n"
    html += "</table>"
    return html


@csrf_exempt
def export_excel(request):
    """
    Экспортирует данные в файл формата Excel, PDF или Word.
    Формат задаётся параметром 'format' в GET-запросе (xlsx, pdf, doc).

    Параметры:
        request: Объект Django HttpRequest, содержащий JSON-данные и параметр 'format'.

    Возвращает:
        HttpResponse с файлом в выбранном формате или сообщение об ошибке.
    """
    try:
        data = json.loads(request.body)
    except json.JSONDecodeError as e:
        return HttpResponse(f'JSONDecodeError: {str(e)}', status=400)

    data = {key: int(value) for key, value in data.items()}
    calculator = Calculator1(data)
    results = calculator.calculate()
    excel = resultToExcel1.fact_plan(results)
    file_format = request.GET.get('format', 'xlsx')

    if file_format == 'xlsx':
        buffer = io.BytesIO()
        excel.save(buffer)
        buffer.seek(0)
        response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename="file.xlsx"'
        return response

    elif file_format == 'pdf':
        pdf = generate_pdf_from_excel(excel)
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="from_excel.pdf"'
        response.write(pdf)
        return response

    elif file_format == 'doc':
        doc = generate_word_from_excel(excel)
        response = HttpResponse(doc, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="file.docx"'
        return response

    return HttpResponse(status=400)


def generate_pdf_from_excel(excel):
    """
    Преобразует данные из Excel в PDF-документ.

    Параметры:
        excel: Объект Excel-книги.

    Возвращает:
        PDF-документ в байтовом формате.
    """
    pdf = FPDF(orientation='L')
    pdf.add_page()
    pdf.add_font('FreeSans', '', 'static/font/FreeSans.ttf', uni=True)
    pdf.set_font('FreeSans', size=7)

    page_width = pdf.w - 20
    column_widths = [
        0.15, 0.1, 0.1, 0.05, 0.05, 0.1, 0.1, 0.1, 0.1, 0.075, 0.075
    ]
    column_widths = [w * page_width for w in column_widths]

    # Печать заголовков
    headers = [
        ' ', 'Факт макс. кол-во файлов', 'Факт разница нагрузки', 'Факт нагрузка в %',
        'Факт нехватка', 'Среднее количество', 'Среднее кол-во файлов',
        'Факт макс. кол-во файлов', 'План разница нагрузки', 'План нагрузка в %', 'План нехватка'
    ]
    header_height = 30
    x_position = 10
    y_position = pdf.get_y()
    for header, width in zip(headers, column_widths):
        pdf.set_xy(x_position, y_position)
        pdf.multi_cell(width, 5, header, border=1, align='C')
        x_position += width
    pdf.ln(header_height)

    # Печать данных
    for row in list(excel.active.iter_rows())[2:]:
        x_position = 10
        for cell, width in zip(row, column_widths):
            value = cell.value
            if value is None:
                for merged_range in excel.active.merged_cells.ranges:
                    if cell.coordinate in merged_range:
                        top_left_cell = excel.active[merged_range.coord.split(":")[0]]
                        value = top_left_cell.value
                        break
            pdf.set_xy(x_position, pdf.get_y())
            pdf.cell(width, 10, str(value) if value else "", border=1, align='C')
            x_position += width
        pdf.ln(10)

    return pdf.output(dest='S').encode('latin-1')


def generate_word_from_excel(excel):
    """
    Преобразует данные из Excel в Word-документ.

    Параметры:
        excel: Объект Excel-книги.

    Возвращает:
        Word-документ в байтовом формате.
    """
    doc = Document()
    doc.add_heading('Экспортированная таблица', level=1)

    sheet = excel.active
    table = doc.add_table(rows=1, cols=sheet.max_column)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    for i, cell in enumerate(sheet[1]):
        hdr_cells[i].text = str(cell.value)

    for row in sheet.iter_rows(min_row=2, values_only=True):
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = str(cell) if cell else ""

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.style.font.size = Pt(10)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc_output = io.BytesIO()
    doc.save(doc_output)
    doc_output.seek(0)
    return doc_output

